"""
Challenge on File merge explorer.
This is the code for searching and merging the files based on the given name.
"""

import datetime
import logging
import os
# import time
import PyPDF2

from tkinter import *
from tkinter import ttk
from zipfile import ZipFile
from docx import Document
from tqdm import tqdm


def getdt():
    return datetime.datetime.now().strftime('%y-%m-%d_%a_%H_%M_%S')  # 21-04-10_Sat_20_53_19


finalmerge = lambda ex: "merge"+ex[1:]+"_"+getdt()+ex
allFilesPath = []
# Initialising the logger file here.
# logging.basicConfig(filename='logtest_gui.log', level=logging.INFO, format='%(levelname)s %(asctime)s %(message)s')
logging.basicConfig(filename='logtest_gui.log',
                    level=logging.INFO,
                    format='%(asctime)s %(levelname)s %(message)s',
                    )


def check_platform():
    """
    This function checks for the OS platform
    :return: 0 for mac, 1 for windows, 2 for linux, otherwise None
    """
    import sys
    if sys.platform == 'darwin':
        logging.info("mac system")
        return 0
    elif sys.platform == 'win32':
        logging.info("Windows system")
        return 1
    elif sys.platform == 'linux':
        logging.info("Linux system")
        return 2
    return None


def getdrives():
    """
    This function is to get the drives present in the system
    :return: drivepaths
    """
    platdct = {0: "mac", 1: "windows", 2: "ubuntu"}
    cp = check_platform()
    if platdct[cp] == 'mac':
        logging.info(os.listdir('/Volumes'))
        available_drives = [os.path.join('/Volumes', i) for i in os.listdir('/Volumes')]
        return available_drives
    elif platdct[cp] == "windows":
        import string
        available_drives = [d + "://" for d in string.ascii_uppercase if os.path.exists(d + "://")]
        logging.info(available_drives)
        return available_drives
    # TODO: for linux
    return None


def getfilepaths(**kwargs):
    """
    This is a generator function which searches for the input string in filepath and generates the filepath
    :param ext: extension to be passed
    :param kwargs: fname --> search string.
    :return: generates filepath
    """
    fname1 = kwargs.get("fname", None)

    if not fname1:
        return []
    fname, ext = '.'.join(fname1.split('.')[:-1]).lower(), fname1.split('.')[-1].lower()
    logging.info(fname + "<==>" + ext)
    all_drives = getdrives()
    for drive in all_drives:
        logging.info("Scanning Drive: " + drive)
        for root, dirs, files in tqdm(os.walk(drive, topdown=True)):
            for name in files[:]:
                if name.lower().find(fname) == -1:  # -1 for didn't match and 0 for match.
                    continue
                # else:
                extn = os.path.splitext(name)[-1].lower()[1:]
                if extn == ext:
                    logging.info("file:" + os.path.join(root, name))
                    yield os.path.join(root, name)


def readfile(ext, filepath):
    """
    This function read the one file and merged file for pdf .
    :param ext:
    :param filepath: filepath to read the file
    :return: reader object
    """
    ret = None
    if ext == '.txt':
        try:
            with open(filepath, 'rb') as infile:
                # with open(filepath, 'r',encoding='utf8') as infile:
                # ret = ""
                mystr = "\n" + filepath + "\n"
                ret = mystr.encode()

                ret += infile.read()
        except Exception as e:
            # PermissionError: [Errno 13]
            # FileNotFoundError: [Errno 2]
            logging.exception(e)
            logging.error(filepath + "FAILED..")
    # elif ext == 'docx':
    #     merged_document = Document()
    #     docxwritepath = finalmerge.get(ext)
    #     if os.path.exists(docxwritepath):
    #         docxs = [filepath, docxwritepath]
    #     else:
    #         docxs = [filepath]
    #     for index, file in enumerate(docxs):
    #         sub_doc = Document(file)
    #
    #         # Don't add a page break if you've reached the last file.
    #         if index < len(docxs) - 1:
    #             sub_doc.add_page_break()
    #
    #         for element in sub_doc.element.body:
    #             merged_document.element.body.append(element)
    #
    #     return merged_document

    elif ext == '.pdf':
        pdfmerger = PyPDF2.PdfFileMerger()
        pdfwritepath = finalmerge(ext)
        if os.path.exists(pdfwritepath):
            pdfs = [filepath, pdfwritepath]
        else:
            pdfs = [filepath]
        for pdf in pdfs:
            try:
                with open(pdf, 'rb') as f:
                    pdfmerger.append(PyPDF2.PdfFileReader(f))
            except Exception as e:
                logging.exception(e)
                logging.error("Failed file " + pdf)

        return pdfmerger
    return ret


def writefile(ext, content_obj, **kwargs):
    """
    This function takes the writer object and saves the file.
    :param ext: extension of the file
    :param content_obj: writer obj
    :param kwargs:
    :return: saved path of the merged files
    """
    ret = finalmerge(ext)
    # if content_obj is None:
    #     print("Content object is None")
    #     return None
    if ext == '.txt':
        try:
            with open(ret, 'ab') as outfile:
                outfile.write(content_obj)
        except Exception as e:
            logging.exception(e)
            logging.error("Merging failed")
    # elif ext == '.docx':
    #     ret = finalmerge[ext]
    #     try:
    #         content_obj.save(ret)
    #     except Exception as e:
    #         logging.exception(e)
    #         logging.error("docx merging failed... ")
    elif ext == '.pdf':
        try:
            with open(ret, 'wb') as outfile:
                content_obj.write(outfile)
        except Exception as e:
            logging.exception(e)
            logging.error("Merging Failed... ")
    return ret


def combine_word_documents(files):
    """
    This function is to merge and save the word documents.
    :param files: docx or doc filepaths
    :return:
    """
    merged_document = Document()

    for index, file in enumerate(files):
        sub_doc = Document(file)

        # Don't add a page break if you've reached the last file.
        if index < len(files) - 1:
            sub_doc.add_page_break()

        for element in sub_doc.element.body:
            merged_document.element.body.append(element)

    merged_document.save(finalmerge('.docx'))


def mergeFiles():
    """
    this function is to merge the files based on the selected extension
    :return:
    """

    extn = cmb.get()
    logging.info(allFilesPath)
    answer.delete(1.0, END)

    fname1 = entry.get()
    fname, ext = '.'.join(fname1.split('.')[:-1]).lower(), fname1.split('.')[-1].lower()
    if extn in ('.docx', '.doc') and ext in ('docx', 'doc'):
        combine_word_documents(allFilesPath)
    elif extn in ('.pdf', '.txt') and ext in ('pdf', 'txt'):
        for filepath in allFilesPath:
            cobj = readfile(extn, filepath)
            writefile(extn, cobj)
    else:
        logging.info("extension not matched., continuing for zipping the files ")
        # fname, ext = '.'.join(fname1.split('.')[:-1]).lower(), fname1.split('.')[-1].lower()
        try:
            if len(allFilesPath) > 0:
                with ZipFile(str(fname1) + "_" + getdt() + ".zip", 'w') as outzipfile:
                    for file in allFilesPath:
                        outzipfile.write(file)
                logging.info("Files zipped and saved here. ")
        except Exception as e:
            logging.error("Failed to zip the files. ")
            logging.exception(e)

    if len(allFilesPath) > 0:
        answer.delete(1.0, END)
        answer.insert(INSERT, f"Merged files successfully saved at {finalmerge(extn)}. ")
        logging.info(f"Merged files successfully saved at {finalmerge(extn)}. ")
        logging.info("Done saving the files. ")
        print("Done saving the files. ")
    else:
        logging.info(f"NO files found to merge for given params {extn} and {entry.get()}")
        answer.insert(INSERT, "NO files found to merge. ")
        print("No files to merge. ")


def search_click(*args):
    logging.info(args)
    fname = entry.get()
    # extn = cmb.get()
    # logging.info("Selected extn: " + cmb.get())
    logging.info("Filename to search: " + fname)
    answer.delete(1.0, END)
    # for idx, filepath in enumerate(getfilepaths(extn, fname=fname)):
    for idx, filepath in enumerate(getfilepaths(fname=fname, search=True)):
        allFilesPath.append(filepath)
        answer.insert(INSERT, str(idx) + "==" + filepath + "\n")
        root.update()
    if len(allFilesPath) > 0:
        answer.insert(INSERT, "...all the files displayed ....")
        # answer.configure(state='disabled')
        logging.info("all Files displayed.... ")
        print("all Files displayed. ")
    else:
        logging.info("NO files found")
        answer.insert(INSERT, "NO files found yet. ")
        print("NO files found yet")


root = Tk()
root.title("Merge Explorer")

topframe = Frame(root)
topframe.pack(side=TOP)
#
# l1 = Label(topframe, text="Select the extension")
# l1.pack()
# extn_options = [".txt", ".pdf", ".docx", ".doc"]
# cmb = ttk.Combobox(topframe, value=extn_options, width=15)
#
# cmb.current(0)
# cmb.pack()
#
# cmb.bind("<<ComboboxSelected>>", search_click)
entry = Entry(topframe)
entry.pack()

button = Button(topframe, text="Search", command=search_click)
button.pack()

bottomframe = Frame(root)
bottomframe.pack(side=BOTTOM)
scroll = Scrollbar(bottomframe)
scroll.pack(side=RIGHT, fill=Y)
answer = Text(bottomframe, yscrollcommand=scroll.set, wrap=WORD)
# answer = Text(bottomframe, width=30, height=10, yscrollcommand=scroll.set, wrap=WORD)
scroll.config(command=answer.yview)
answer.pack()
# if True:

l1 = Label(bottomframe, text="Select the extension")
l1.pack()
extn_options = [".txt", ".pdf", ".docx", ".doc", 'other']
cmb = ttk.Combobox(bottomframe, value=extn_options, width=15)

cmb.current(4)
cmb.pack()

# cmb.bind("<<ComboboxSelected>>", mergeFiles)
button = Button(bottomframe, text="Merge", command=mergeFiles)
button.pack()

root.configure(background='ivory3')
root.mainloop()
# Screenshot 2021-03-30.png
